from unittest.mock import MagicMock, patch

from docx import Document

from core.cv_tailor import (
    append_skill,
    apply_edit,
    apply_proposals,
    find_category_cell,
    find_skills_table,
    is_editable,
    pick_target_category,
    propose_edits,
)


def _uniform_paragraph():
    """A paragraph split into two runs that share identical formatting.

    Word does this routinely (spell-check/autocorrect boundaries) even
    with no visual difference between the runs.
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Senior Data ")
    p.add_run("Engineer")
    return p


def _mixed_paragraph():
    """A bold title followed by a plain date, on one line — genuinely mixed."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Senior Data Engineer").bold = True
    p.add_run(" (2019-2023)")
    return p


def test_is_editable_true_for_uniform_formatting():
    assert is_editable(_uniform_paragraph()) is True


def test_is_editable_false_for_mixed_formatting():
    assert is_editable(_mixed_paragraph()) is False


def test_is_editable_false_for_empty_paragraph():
    doc = Document()
    p = doc.add_paragraph()
    assert is_editable(p) is False


def test_apply_edit_replaces_text_and_preserves_formatting():
    p = _uniform_paragraph()
    apply_edit(p, "Lead Data Engineer")
    assert p.text == "Lead Data Engineer"
    assert p.runs[0].text == "Lead Data Engineer"
    assert all(run.text == "" for run in p.runs[1:])


def test_apply_edit_raises_on_mixed_formatting():
    p = _mixed_paragraph()
    try:
        apply_edit(p, "won't be applied")
        assert False, "expected ValueError"
    except ValueError:
        pass
    assert p.text == "Senior Data Engineer (2019-2023)"


def _skills_table():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Programming")
    table.rows[0].cells[1].paragraphs[0].add_run("Python, SQL")
    table.rows[1].cells[0].paragraphs[0].add_run("BI and Reporting")
    table.rows[1].cells[1].paragraphs[0].add_run("Power BI")
    return table


def test_find_category_cell_returns_matching_cell():
    table = _skills_table()
    cell = find_category_cell(table, "BI and Reporting")
    assert cell.text == "Power BI"


def test_find_category_cell_returns_none_when_not_found():
    table = _skills_table()
    assert find_category_cell(table, "Nonexistent Category") is None


def test_append_skill_adds_comma_separated_when_existing_text():
    table = _skills_table()
    cell = find_category_cell(table, "Programming")
    append_skill(cell, "Rust")
    assert cell.text == "Python, SQL, Rust"


def test_append_skill_raises_when_cell_not_editable():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    cell = table.rows[0].cells[1]
    p = cell.paragraphs[0]
    p.add_run("Python").bold = True
    p.add_run(", SQL")
    try:
        append_skill(cell, "Rust")
        assert False, "expected ValueError"
    except ValueError:
        pass
    assert cell.text == "Python, SQL"


def test_find_skills_table_matches_category_shaped_table():
    doc = Document()
    doc.add_table(rows=1, cols=1)  # wrong shape, e.g. a single-column layout table
    skills = doc.add_table(rows=2, cols=2)
    skills.rows[0].cells[0].paragraphs[0].add_run("Programming")
    skills.rows[1].cells[0].paragraphs[0].add_run("BI and Reporting")
    table = find_skills_table(doc)
    assert table is not None
    assert table.rows[0].cells[0].text == "Programming"


def test_find_skills_table_skips_contact_info_table():
    doc = Document()
    contact = doc.add_table(rows=1, cols=2)
    contact.rows[0].cells[0].paragraphs[0].add_run(
        "Hawthorn, VIC 3122\nerfan.hesami@outlook.com\n+61 406 250 800"
    )
    assert find_skills_table(doc) is None


def test_find_skills_table_returns_none_when_no_table_present():
    doc = Document()
    assert find_skills_table(doc) is None


def _skills_doc():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Programming")
    table.rows[0].cells[1].paragraphs[0].add_run("Python, SQL")
    table.rows[1].cells[0].paragraphs[0].add_run("BI and Reporting")
    table.rows[1].cells[1].paragraphs[0].add_run("Power BI")
    return doc


@patch("core.cv_tailor.tracked_call")
def test_propose_edits_returns_proposal_for_buried_skill(mock_call):
    mock_call.return_value = _fake_llm_response('{"category": "Programming"}')
    doc = _skills_doc()
    gap_suggestions = [{"skill": "Rust", "status": "buried", "evidence": "Wrote Rust CLI tools"}]
    proposals = propose_edits(doc, gap_suggestions)
    assert proposals == [{
        "skills": [{"skill": "Rust", "evidence": "Wrote Rust CLI tools"}],
        "category": "Programming",
        "current_text": "Python, SQL",
        "proposed_text": "Python, SQL, Rust",
    }]


@patch("core.cv_tailor.tracked_call")
def test_propose_edits_groups_same_category_skills_into_one_proposal(mock_call):
    # Both skills resolve to "Programming" — must become ONE combined proposal,
    # not two independent ones that would clobber each other when applied.
    mock_call.return_value = _fake_llm_response('{"category": "Programming"}')
    doc = _skills_doc()
    gap_suggestions = [
        {"skill": "Rust", "status": "buried", "evidence": "Wrote Rust CLI tools"},
        {"skill": "Go", "status": "buried", "evidence": "Wrote Go microservices"},
    ]
    proposals = propose_edits(doc, gap_suggestions)
    assert len(proposals) == 1
    assert proposals[0]["proposed_text"] == "Python, SQL, Rust, Go"
    assert proposals[0]["skills"] == [
        {"skill": "Rust", "evidence": "Wrote Rust CLI tools"},
        {"skill": "Go", "evidence": "Wrote Go microservices"},
    ]


@patch("core.cv_tailor.tracked_call")
def test_propose_edits_skips_missing_status(mock_call):
    doc = _skills_doc()
    gap_suggestions = [{"skill": "Docker", "status": "missing", "evidence": None}]
    assert propose_edits(doc, gap_suggestions) == []
    mock_call.assert_not_called()


@patch("core.cv_tailor.tracked_call")
def test_propose_edits_skips_when_no_category_fits(mock_call):
    mock_call.return_value = _fake_llm_response('{"category": null}')
    doc = _skills_doc()
    gap_suggestions = [{"skill": "Interpretive Dance", "status": "buried", "evidence": "some evidence"}]
    assert propose_edits(doc, gap_suggestions) == []


def test_propose_edits_returns_empty_when_no_skills_table():
    doc = Document()  # no tables at all
    gap_suggestions = [{"skill": "Rust", "status": "buried", "evidence": "x"}]
    assert propose_edits(doc, gap_suggestions) == []


def test_apply_proposals_applies_possibly_edited_text():
    doc = _skills_doc()
    proposals = [{
        "skill": "Rust",
        "evidence": "Wrote Rust CLI tools",
        "category": "Programming",
        "current_text": "Python, SQL",
        "proposed_text": "Python, SQL, Rust (systems programming)",  # user edited this in the review UI
    }]
    apply_proposals(doc, proposals)
    cell = find_category_cell(find_skills_table(doc), "Programming")
    assert cell.text == "Python, SQL, Rust (systems programming)"


def _fake_llm_response(json_text):
    message = MagicMock()
    message.content = [MagicMock(text=json_text)]
    return message


CATEGORIES = ["Data Platforms and Warehousing", "Programming", "BI and Reporting"]


@patch("core.cv_tailor.tracked_call")
def test_pick_target_category_returns_valid_category(mock_call):
    mock_call.return_value = _fake_llm_response('{"category": "Programming"}')
    result = pick_target_category(CATEGORIES, "Rust", "Wrote Rust CLI tools")
    assert result == "Programming"


@patch("core.cv_tailor.tracked_call")
def test_pick_target_category_returns_none_when_model_says_null(mock_call):
    mock_call.return_value = _fake_llm_response('{"category": null}')
    result = pick_target_category(CATEGORIES, "Interpretive Dance", "no real fit")
    assert result is None


@patch("core.cv_tailor.tracked_call")
def test_pick_target_category_returns_none_when_model_hallucinates_category(mock_call):
    # Model returns a category that wasn't in the list we gave it — the code
    # guard must reject this rather than trust it, same honesty-line pattern
    # used in matcher.py.
    mock_call.return_value = _fake_llm_response('{"category": "Made Up Category"}')
    result = pick_target_category(CATEGORIES, "Rust", "Wrote Rust CLI tools")
    assert result is None
